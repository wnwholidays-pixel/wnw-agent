import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import os
import shutil
import subprocess
import tempfile


# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="WNW Template Engine",
    layout="wide"
)

st.title("🦅 Wings 'N' Wheels Holidays")
st.caption(
    "Official Production Studio - Final Master Template Merger Engine"
)


# ============================================================
# SIDEBAR - BOOKING PROFILE
# ============================================================

with st.sidebar:

    st.header("📋 Booking Profile")

    school_name = st.text_input(
        "School Name:",
        "AA School"
    )

    start_city = st.text_input(
        "Starting From:",
        "Jaipur"
    )

    destination_name = st.text_input(
        "Destination Label:",
        "CHANDIGARH – MANALI"
    )

    tour_duration = st.text_input(
        "Duration Frame:",
        "6 Nights / 7 Days"
    )

    st.header("💰 Pricing Matrix")

    student_cost = st.text_input(
        "Cost Per Student:",
        "Rs. 13,500/-"
    )

    group_strength = st.text_input(
        "Group Strength:",
        "45 (Minimum)"
    )

    teacher_ratio = st.text_input(
        "Teacher Ratio:",
        "15:01"
    )


# ============================================================
# ITINERARY INPUT
# ============================================================

pasted_itinerary = st.text_area(
    "Pasted Itinerary Body Text:",
    height=600,
    placeholder=(
        "Paste the complete formatted itinerary here "
        "starting from TABLE_START..."
    )
)


# ============================================================
# STYLE CONSTANTS
# ============================================================

BLUE_DAY = RGBColor(0, 163, 224)
BLUE_ROUTE = RGBColor(0, 86, 179)
GREY_DIVIDER = RGBColor(100, 100, 100)

FONT_NAME = "Arial"


# ============================================================
# DAY HEADING FORMATTER
# ============================================================

def format_day_heading(paragraph, day_number):

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        f"DAY {day_number}"
    )

    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BLUE_DAY


# ============================================================
# ROUTE FORMATTER
# ============================================================

def format_route(paragraph, route_text):

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(
        route_text
    )

    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BLUE_ROUTE


# ============================================================
# TIMELINE LINE FORMATTER
# ============================================================

def format_timeline_line(paragraph, line):

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Detect 12-hour timestamp
    time_match = re.match(
        r"^(\d{1,2}:\d{2}\s?(?:AM|PM))\s+(.*)$",
        line.strip(),
        re.IGNORECASE
    )

    if time_match:

        time_text = time_match.group(1)
        event_text = time_match.group(2)

        # Bold timestamp
        time_run = paragraph.add_run(
            time_text + "\t"
        )

        time_run.font.name = FONT_NAME
        time_run.font.size = Pt(11)
        time_run.font.bold = True

        # Normal event text
        event_run = paragraph.add_run(
            event_text
        )

        event_run.font.name = FONT_NAME
        event_run.font.size = Pt(11)

    else:

        run = paragraph.add_run(
            line.strip()
        )

        run.font.name = FONT_NAME
        run.font.size = Pt(11)


# ============================================================
# ADD DETAILED ITINERARY LINE
# ============================================================

def append_styled_line(doc, current_paragraph, line):

    stripped = line.strip()

    if not stripped:
        return current_paragraph

    new_paragraph = doc.add_paragraph()


    # ========================================================
    # DAY HEADING
    # ========================================================

    if (
        stripped.upper().startswith("DAY ")
        and ":" in stripped
    ):

        numbers = re.findall(
            r"\d+",
            stripped
        )

        day_number = (
            numbers[0]
            if numbers
            else "1"
        )

        format_day_heading(
            new_paragraph,
            day_number
        )


    # ========================================================
    # SUB ROUTE
    # ========================================================

    elif stripped.upper().startswith("[SUB_ROUTE]"):

        route_text = re.sub(
            r"^\[SUB_ROUTE\]\s*",
            "",
            stripped,
            flags=re.IGNORECASE
        )

        format_route(
            new_paragraph,
            route_text
        )


    # ========================================================
    # OVERNIGHT DIVIDER
    # ========================================================

    elif (
        "Overnight Journey" in stripped
        or "---" in stripped
    ):

        new_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = new_paragraph.add_run(
            stripped
        )

        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = GREY_DIVIDER


    # ========================================================
    # NORMAL / TIMESTAMP EVENT
    # ========================================================

    else:

        format_timeline_line(
            new_paragraph,
            stripped
        )


    # ========================================================
    # INSERT AFTER CURRENT PARAGRAPH
    # ========================================================

    current_paragraph._p.addnext(
        new_paragraph._p
    )

    return new_paragraph


# ============================================================
# FIND LIBREOFFICE
# ============================================================

def find_libreoffice():

    possible_commands = [

        # Linux / Streamlit Cloud
        "libreoffice",

        # Alternative Linux executable
        "soffice",

        # Windows installation
        r"C:\Program Files\LibreOffice\program\soffice.exe",

        # Windows 32-bit installation
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",

        # User installation
        os.path.expanduser(
            r"~\AppData\Local\Programs\LibreOffice\program\soffice.exe"
        )
    ]


    for command in possible_commands:

        # Full path exists
        if os.path.isfile(command):

            return command

        # Command exists in PATH
        if shutil.which(command):

            return shutil.which(command)


    return None


# ============================================================
# CONVERT WORD → PDF
# ============================================================

def convert_docx_to_pdf(docx_bytes):

    libreoffice = find_libreoffice()

    if libreoffice is None:

        raise RuntimeError(
            "LibreOffice was not found. "
            "Please install LibreOffice or add it to "
            "the Streamlit deployment."
        )


    with tempfile.TemporaryDirectory() as temp_dir:

        docx_path = os.path.join(
            temp_dir,
            "WNW_Itinerary.docx"
        )

        pdf_path = os.path.join(
            temp_dir,
            "WNW_Itinerary.pdf"
        )


        # ----------------------------------------------------
        # SAVE TEMPORARY WORD FILE
        # ----------------------------------------------------

        with open(
            docx_path,
            "wb"
        ) as file:

            file.write(
                docx_bytes
            )


        # ----------------------------------------------------
        # CONVERT TO PDF
        # ----------------------------------------------------

        process = subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                temp_dir,
                docx_path
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )


        # ----------------------------------------------------
        # CHECK PROCESS
        # ----------------------------------------------------

        if process.returncode != 0:

            error_message = (
                process.stderr.strip()
                or process.stdout.strip()
                or "LibreOffice conversion failed."
            )

            raise RuntimeError(
                error_message
            )


        # ----------------------------------------------------
        # CHECK PDF
        # ----------------------------------------------------

        if not os.path.exists(pdf_path):

            raise RuntimeError(
                "LibreOffice completed, but "
                "the PDF file was not created."
            )


        # ----------------------------------------------------
        # READ PDF
        # ----------------------------------------------------

        with open(
            pdf_path,
            "rb"
        ) as file:

            return file.read()


# ============================================================
# PARSE ITINERARY
# ============================================================

def parse_itinerary(text):

    lines = text.splitlines()

    table_rows = []
    detailed_lines = []
    inclusions = []
    exclusions = []

    current_mode = "detailed"


    for raw_line in lines:

        line = raw_line.strip()


        # ----------------------------------------------------
        # BLOCK START / END
        # ----------------------------------------------------

        if "TABLE_START" in line:

            current_mode = "table"
            continue


        if "TABLE_END" in line:

            current_mode = "detailed"
            continue


        if "INCLUSIONS_START" in line:

            current_mode = "inclusions"
            continue


        if "INCLUSIONS_END" in line:

            current_mode = "detailed"
            continue


        if "EXCLUSIONS_START" in line:

            current_mode = "exclusions"
            continue


        if "EXCLUSIONS_END" in line:

            current_mode = "detailed"
            continue


        # ----------------------------------------------------
        # SUMMARY TABLE
        # ----------------------------------------------------

        if current_mode == "table":

            if "|" in line:

                cells = [
                    cell.strip()
                    for cell in line.split("|")
                    if cell.strip()
                ]


                # Ignore table header
                if (
                    len(cells) >= 2
                    and cells[0].lower() != "day"
                ):

                    table_rows.append(
                        cells
                    )


        # ----------------------------------------------------
        # INCLUSIONS
        # ----------------------------------------------------

        elif current_mode == "inclusions":

            if line:

                inclusions.append(
                    line
                )


        # ----------------------------------------------------
        # EXCLUSIONS
        # ----------------------------------------------------

        elif current_mode == "exclusions":

            if line:

                exclusions.append(
                    line
                )


        # ----------------------------------------------------
        # DETAILED ITINERARY
        # ----------------------------------------------------

        else:

            if line:

                detailed_lines.append(
                    line
                )


    return (
        table_rows,
        detailed_lines,
        inclusions,
        exclusions
    )


# ============================================================
# CALCULATE TOUR ROUTE
# ============================================================

def calculate_tour_route(
    table_rows,
    start_city
):

    route_cities = []


    for row in table_rows:

        if len(row) <= 1:
            continue


        city_field = str(
            row[1]
        ).upper()


        cleaned = (
            city_field
            .replace("[", "")
            .replace("]", "")
            .replace("'", "")
            .replace('"', "")
        )


        cities = [
            city.strip()
            for city in cleaned.split("→")
            if city.strip()
        ]


        for city in cities:

            if city not in route_cities:

                route_cities.append(
                    city
                )


    if (
        route_cities
        and start_city.upper()
        not in route_cities[-1]
    ):

        route_cities.append(
            start_city.upper()
        )


    return " → ".join(
        route_cities
    )


# ============================================================
# PROCESS TABLE ROWS
# ============================================================

def process_table_rows(table_rows):

    processed = []


    for row in table_rows:

        if len(row) >= 5:

            # First four columns remain unchanged
            # Everything after column 4 = meals
            combined_meals = " | ".join(
                row[4:]
            )

            processed.append(
                row[:4] + [combined_meals]
            )

        else:

            processed.append(
                row
            )


    return processed


# ============================================================
# MAIN COMPILE BUTTON
# ============================================================

if st.button(
    "🚀 Compile Official Word + PDF Proposal",
    type="primary"
):

    # ========================================================
    # CHECK INPUT
    # ========================================================

    if not pasted_itinerary.strip():

        st.error(
            "Please paste the formatted itinerary text block first!"
        )

        st.stop()


    try:

        # ====================================================
        # LOAD MASTER TEMPLATE
        # ====================================================

        doc = Document(
            "template.docx"
        )


        # ====================================================
        # PARSE INPUT
        # ====================================================

        (
            table_rows_data,
            detailed_lines,
            inclusions,
            exclusions
        ) = parse_itinerary(
            pasted_itinerary
        )


        # ====================================================
        # CALCULATE ROUTE
        # ====================================================

        calculated_tour_route = calculate_tour_route(
            table_rows_data,
            start_city
        )


        # ====================================================
        # PROCESS TABLE
        # ====================================================

        processed_table_rows = process_table_rows(
            table_rows_data
        )


        # ====================================================
        # REPLACE TEMPLATE VARIABLES
        # ====================================================

        for paragraph in doc.paragraphs:


            # ------------------------------------------------
            # SCHOOL NAME
            # ------------------------------------------------

            if "{{SCHOOL_NAME}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{SCHOOL_NAME}}",
                    school_name
                )


            # ------------------------------------------------
            # DESTINATION
            # ------------------------------------------------

            if "{{DESTINATION_NAME}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{DESTINATION_NAME}}",
                    destination_name.upper()
                )


            # ------------------------------------------------
            # TOUR DURATION
            # ------------------------------------------------

            if "{{TOUR_DURATION}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{TOUR_DURATION}}",
                    tour_duration
                )


            # ------------------------------------------------
            # TOUR ROUTE
            # ------------------------------------------------

            if "{{TOUR_ROUTE}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{TOUR_ROUTE}}",
                    calculated_tour_route
                )


            # =================================================
            # INCLUSIONS
            # =================================================

            if "{{TOUR_INCLUSIONS}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{TOUR_INCLUSIONS}}",
                    ""
                )

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                current = paragraph


                for item in inclusions:

                    clean_item = re.sub(
                        r"^•\s*",
                        "",
                        item
                    )


                    new_paragraph = doc.add_paragraph(
                        clean_item,
                        style="List Bullet"
                    )


                    new_paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                    )


                    for run in new_paragraph.runs:

                        run.font.name = FONT_NAME
                        run.font.size = Pt(10.5)


                    current._p.addnext(
                        new_paragraph._p
                    )

                    current = new_paragraph


            # =================================================
            # EXCLUSIONS
            # =================================================

            if "{{TOUR_EXCLUSIONS}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{TOUR_EXCLUSIONS}}",
                    ""
                )

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                current = paragraph


                for item in exclusions:

                    clean_item = re.sub(
                        r"^•\s*",
                        "",
                        item
                    )


                    new_paragraph = doc.add_paragraph(
                        clean_item,
                        style="List Bullet"
                    )


                    new_paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                    )


                    for run in new_paragraph.runs:

                        run.font.name = FONT_NAME
                        run.font.size = Pt(10.5)


                    current._p.addnext(
                        new_paragraph._p
                    )

                    current = new_paragraph


            # =================================================
            # DETAILED ITINERARY
            # =================================================

            if "{{DETAILED_ITINERARY}}" in paragraph.text:

                paragraph.text = paragraph.text.replace(
                    "{{DETAILED_ITINERARY}}",
                    ""
                )

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                current = paragraph


                for line in detailed_lines:

                    current = append_styled_line(
                        doc,
                        current,
                        line
                    )


        # ====================================================
        # PROCESS ALL TABLES
        # ====================================================

        for table in doc.tables:

            target_table = False
            target_row_index = -1


            # ------------------------------------------------
            # FIND SUMMARY TABLE
            # ------------------------------------------------

            for row_index, row in enumerate(
                table.rows
            ):

                for cell in row.cells:

                    if "{{DAY_NUM}}" in cell.text:

                        target_table = True
                        target_row_index = row_index

                        break


                if target_table:
                    break


            # ------------------------------------------------
            # INSERT SUMMARY TABLE ROWS
            # ------------------------------------------------

            if (
                target_table
                and target_row_index >= 0
            ):

                for index, row_data in enumerate(
                    processed_table_rows
                ):


                    # First row uses template row
                    if index == 0:

                        new_row = table.rows[
                            target_row_index
                        ]

                    else:

                        new_row = table.add_row()


                    # ----------------------------------------
                    # WRITE CELLS
                    # ----------------------------------------

                    for cell_index in range(
                        min(
                            len(row_data),
                            len(new_row.cells)
                        )
                    ):

                        cell_text = row_data[
                            cell_index
                        ]


                        # ------------------------------------
                        # STACK MEALS VERTICALLY
                        # ------------------------------------

                        if cell_index == 4:

                            cell_text = (
                                cell_text
                                .replace(
                                    "B:",
                                    "B:"
                                )
                                .replace(
                                    "L:",
                                    "\nL:"
                                )
                                .replace(
                                    "D:",
                                    "\nD:"
                                )
                            )


                        new_row.cells[
                            cell_index
                        ].text = cell_text


                        # ------------------------------------
                        # TABLE FONT
                        # ------------------------------------

                        for p in new_row.cells[
                            cell_index
                        ].paragraphs:

                            for run in p.runs:

                                run.font.name = FONT_NAME
                                run.font.size = Pt(9)


            # =================================================
            # PRICING VARIABLES
            # =================================================

            for row in table.rows:

                for cell in row.cells:


                    # -----------------------------------------
                    # STUDENT COST
                    # -----------------------------------------

                    if "{{STUDENT_COST}}" in cell.text:

                        cell.text = cell.text.replace(
                            "{{STUDENT_COST}}",
                            ""
                        )


                        run = (
                            cell.paragraphs[0]
                            .add_run(student_cost)
                        )


                        run.font.name = FONT_NAME
                        run.font.size = Pt(14)
                        run.font.bold = True


                    # -----------------------------------------
                    # GROUP STRENGTH
                    # -----------------------------------------

                    if "{{GROUP_STRENGTH}}" in cell.text:

                        cell.text = cell.text.replace(
                            "{{GROUP_STRENGTH}}",
                            group_strength
                        )


                    # -----------------------------------------
                    # TEACHER RATIO
                    # -----------------------------------------

                    if "{{TEACHER_RATIO}}" in cell.text:

                        cell.text = cell.text.replace(
                            "{{TEACHER_RATIO}}",
                            teacher_ratio
                        )


        # ====================================================
        # SAVE WORD DOCUMENT
        # ====================================================

        word_buffer = io.BytesIO()

        doc.save(
            word_buffer
        )

        word_data = word_buffer.getvalue()


        # ====================================================
        # WORD SUCCESS
        # ====================================================

        st.success(
            "🎉 Final Word document compiled successfully!"
        )


        # ====================================================
        # CREATE PDF
        # ====================================================

        pdf_data = None

        try:

            pdf_data = convert_docx_to_pdf(
                word_data
            )

            st.success(
                "📄 Final PDF document created successfully!"
            )


        except Exception as pdf_error:

            st.warning(
                "Word document was created successfully, "
                "but PDF conversion failed."
            )

            st.error(
                f"PDF conversion error: {str(pdf_error)}"
            )


        # ====================================================
        # DOWNLOAD BUTTONS
        # ====================================================

        file_name_base = (
            "WNW_Itinerary_"
            + school_name.replace(" ", "_")
        )


        col1, col2 = st.columns(2)


        # ====================================================
        # WORD DOWNLOAD
        # ====================================================

        with col1:

            st.download_button(
                label="💾 Download Word Document (.docx)",
                data=word_data,
                file_name=(
                    file_name_base
                    + ".docx"
                ),
                mime=(
                    "application/"
                    "vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
            )


        # ====================================================
        # PDF DOWNLOAD
        # ====================================================

        with col2:

            if pdf_data is not None:

                st.download_button(
                    label="📄 Download PDF Document (.pdf)",
                    data=pdf_data,
                    file_name=(
                        file_name_base
                        + ".pdf"
                    ),
                    mime="application/pdf"
                )

            else:

                st.button(
                    "📄 PDF Unavailable",
                    disabled=True
                )


    # ========================================================
    # MAIN ERROR HANDLER
    # ========================================================

    except FileNotFoundError:

        st.error(
            "❌ template.docx was not found. "
            "Please keep template.docx in the same folder "
            "as your Streamlit app."
        )


    except Exception as error:

        st.error(
            "❌ Error merging template data: "
            + str(error)
        )
